# -*- coding: utf-8 -*-
"""
report.py — Genera el reporte astral interpretativo en PDF y DOCX.

Estructura del documento:
  Portada (azul noche · tipografía serif · datos completos de la carta)
  Índice de contenidos (con números de página y enlaces internos)
  1. Qué es una carta astral natal (+ mandamientos, cruz, elemento)
  2. Carta natal: el libreto de tu vida (rueda + leyenda de símbolos)
  3. Los planetas en tu carta natal
  4. Las cúspides de las casas
  5. Aspectos astrales
  6. Glosario · Comentarios finales
  Redes sociales del astrólogo

Textos: interpretations_es.json + preamble_es.json + glossary_es.json + brand_es.json.
"""
import os
import io
import json
import base64
from datetime import date as _date

_BASE = os.path.dirname(os.path.abspath(__file__))

try:
    import report_chart as _rc          # lámina estética de carta (solo reportes)
except Exception:
    _rc = None

GOLD = '#C9A24B'
NAVY = '#1E2450'
BLUE = '#3A4488'
SLATE = '#8A93B5'
INK = '#22284A'
RED = '#B02020'
PAPER_BG = '#FFFFFF'
LIGHT = '#E8ECF8'

ROMAN = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI', 'XII']

_NAME_BY_KEY = {
    'aSol': 'Sol', 'aLuna': 'Luna', 'aMercurio': 'Mercurio', 'aVenus': 'Venus',
    'aMarte': 'Marte', 'aJupiter': 'Júpiter', 'aSaturno': 'Saturno',
    'aUrano': 'Urano', 'aNeptuno': 'Neptuno', 'aPluton': 'Plutón',
    'aChiron': 'Quirón', 'aNoduloNorte': 'Nodo Norte', 'aNoduloSur': 'Nodo Sur',
    'aLunaNegra': 'Luna Negra', 'aRuedaFortuna': 'Parte de la Fortuna',
}
MESES = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio', 'julio',
         'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
MESES_AB = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun', 'Jul', 'Ago', 'Sep',
            'Oct', 'Nov', 'Dic']


def fecha_es(iso):
    """'1988-09-28' → '28 de septiembre de 1988'."""
    try:
        p = iso.replace('/', '-').split('-')
        return "%d de %s de %s" % (int(p[2]), MESES[int(p[1]) - 1], p[0])
    except Exception:
        return iso


def fecha_ab(iso):
    """'1988-09-28' → 'Sep 28. 1988'."""
    try:
        p = iso.replace('/', '-').split('-')
        return "%s %d. %s" % (MESES_AB[int(p[1]) - 1], int(p[2]), p[0])
    except Exception:
        return iso


# ── Íconos propios ──────────────────────────────────────────────────────
_ICONS_DIR = os.path.join(_BASE, 'icons')

# ── Portadas de marca y página de información (imágenes de página completa) ─
_COVERS_DIR = os.path.join(_BASE, 'covers')
_COVER_FILE = {'natal': 'natal.jpg', 'solar_return': 'solar_return.jpg',
               'combined': 'combined.jpg', 'akashic': 'akashic.jpg'}


def _cover_path(chart_type):
    f = _COVER_FILE.get(chart_type)
    if f:
        p = os.path.join(_COVERS_DIR, f)
        if os.path.exists(p):
            return p
    return None


def _info_path():
    p = os.path.join(_COVERS_DIR, 'info.jpg')
    return p if os.path.exists(p) else None


# Aviso legal y de privacidad (página penúltima del reporte)
_DISCLAIMER = [
    ("Tu privacidad",
     "En https://www.mauriciopuerta.tv/ nos comprometemos a respetar tu "
     "privacidad. Reconocemos que cuando eliges proporcionar información, "
     "confías en que actuaremos de manera responsable. Es por eso que hemos "
     "establecido una política para proteger tu información personal. Siempre "
     "puedes visitar y navegar el sitio web sin necesidad de darnos tu "
     "información personal."),
    ("Nota legal",
     "Si bien la astrología no es una ciencia exacta, te recomendamos "
     "encarecidamente no tomar literalmente las interpretaciones que puedas "
     "recibir con una interpretación. Utiliza tu sentido común y tu propio "
     "juicio; las interpretaciones están destinadas a reflejar procesos propios "
     "de pensamiento. La astrología brinda información, mas no predice "
     "absolutamente los eventos particulares de tus decisiones, ni te da "
     "consejos absolutos sobre las acciones o las decisiones que debes tomar "
     "con respecto a las circunstancias de tu vida actual."),
    ("Oficial y legal",
     "https://www.mauriciopuerta.tv, Saturn Marketing LLC o Mauricio Puerta® no "
     "hacen representaciones o garantías de ningún tipo, expresas o implícitas, "
     "en cuanto a la operación del sitio, información, contenido, materiales o "
     "productos incluidos. https://www.mauriciopuerta.tv, Saturn Marketing LLC o "
     "Mauricio Puerta® no serán responsables de daños derivados del uso de esta "
     "información. Todas las interpretaciones y consejos derivados del uso de "
     "este sitio se entienden solo para fines de entretenimiento."),
]

_NAME_ICON = [
    ("Luna Negra", "aLunaNegra"), ("Nodo Norte", "aNoduloNorte"),
    ("Nodo Sur", "aNoduloSur"), ("Nódulo Lunar Norte", "aNoduloNorte"),
    ("Nódulo Lunar Sur", "aNoduloSur"), ("Nódulo Norte", "aNoduloNorte"),
    ("Nódulo Sur", "aNoduloSur"), ("Rueda de la Fortuna", "aRuedaFortuna"),
    ("Parte de la Fortuna", "aRuedaFortuna"),
    ("Sol", "aSol"), ("Luna", "aLuna"), ("Mercurio", "aMercurio"),
    ("Venus", "aVenus"), ("Marte", "aMarte"), ("Júpiter", "aJupiter"),
    ("Jupiter", "aJupiter"), ("Saturno", "aSaturno"), ("Urano", "aUrano"),
    ("Neptuno", "aNeptuno"), ("Plutón", "aPluton"), ("Pluton", "aPluton"),
    ("Quirón", "aChiron"), ("Quiron", "aChiron"), ("Tierra", "aTierra"),
    ("Aries", "Aries"), ("Tauro", "Taurus"), ("Géminis", "Gemini"),
    ("Geminis", "Gemini"), ("Cáncer", "Cancer"), ("Cancer", "Cancer"),
    ("Leo", "Leo"), ("Virgo", "Virgo"), ("Libra", "Libra"),
    ("Escorpión", "Scorpio"), ("Escorpio", "Scorpio"),
    ("Sagitario", "Sagittarius"), ("Capricornio", "Capricorn"),
    ("Acuario", "Aquarius"), ("Piscis", "Pisces"),
]

import re as _re
_NAME_RE = _re.compile(
    r'\b(' + '|'.join(_re.escape(n) for n, _ in _NAME_ICON) + r')\b')
_NAME2FILE = {n: k for n, k in _NAME_ICON}


def _icon_path(name):
    p = os.path.join(_ICONS_DIR, _NAME2FILE.get(name, '') + '.png')
    return p if os.path.exists(p) else None


def _segments(text):
    out = []
    last = 0
    for m in _NAME_RE.finditer(text):
        out.append((text[last:m.end()], _icon_path(m.group(1))))
        last = m.end()
    out.append((text[last:], None))
    return out


# ── Datos ───────────────────────────────────────────────────────────────
_CACHE = {}


def _load_json(fname):
    if fname not in _CACHE:
        try:
            with open(os.path.join(_BASE, fname), encoding='utf-8') as f:
                _CACHE[fname] = json.load(f)
        except Exception:
            _CACHE[fname] = {}
    return _CACHE[fname]


def _interp():
    return _load_json('interpretations_es.json')


def _preamble():
    return _load_json('preamble_es.json')


def _glossary():
    return _load_json('glossary_es.json')


def _brand():
    return _load_json('brand_es.json')


def _akashic():
    return _load_json('akashic_es.json')


_REGENTE = {"Aries": "Marte", "Taurus": "Venus", "Gemini": "Mercurio",
            "Cancer": "Luna", "Leo": "Sol", "Virgo": "Quirón", "Libra": "Venus",
            "Scorpio": "Plutón", "Sagittarius": "Júpiter", "Capricorn": "Saturno",
            "Aquarius": "Urano", "Pisces": "Neptuno"}
_SIGNS_EN = ['Aries', 'Taurus', 'Gemini', 'Cancer', 'Leo', 'Virgo', 'Libra',
             'Scorpio', 'Sagittarius', 'Capricorn', 'Aquarius', 'Pisces']
_SIGNS_ES = ['Aries', 'Tauro', 'Géminis', 'Cáncer', 'Leo', 'Virgo', 'Libra',
             'Escorpio', 'Sagitario', 'Capricornio', 'Acuario', 'Piscis']
_ES_BY_EN = dict(zip(_SIGNS_EN, _SIGNS_ES))

# Regentes de decanatos por triplicidad (mismo sistema de la rueda, con Quirón)
_DECAN_RULERS = [
    ["aMarte", "aSol", "aJupiter"],    ["aVenus", "aChiron", "aSaturno"],
    ["aMercurio", "aVenus", "aUrano"], ["aLuna", "aPluton", "aNeptuno"],
    ["aSol", "aJupiter", "aMarte"],    ["aChiron", "aSaturno", "aVenus"],
    ["aVenus", "aUrano", "aMercurio"], ["aPluton", "aNeptuno", "aLuna"],
    ["aJupiter", "aMarte", "aSol"],    ["aSaturno", "aVenus", "aChiron"],
    ["aUrano", "aMercurio", "aVenus"], ["aNeptuno", "aLuna", "aPluton"]]
_ORD_DECAN = ["primer", "segundo", "tercer"]


def build_decanates(chart):
    """Para cada casa: en qué decanato cae su cúspide, quién rige ese decanato
    y dónde está ese regente (signo y casa). Relaciona decanato → regente →
    su ubicación, como pide el método. Aparece en todos los reportes."""
    planets = {p['key']: p for p in chart['planets']}
    items = []
    for h in chart.get('houses', []):
        si = h.get('sign_index', 1) - 1
        deg = h.get('deg', 0)
        dec = min(2, int(deg) // 10)
        try:
            rk = _DECAN_RULERS[si][dec]
        except Exception:
            continue
        reg = planets.get(rk)
        reg_name = _NAME_BY_KEY.get(rk, rk)
        title = "Casa %s · cúspide en %s" % (h.get('roman', ''), h.get('sign_es', h.get('sign', '')))
        if reg:
            reg_sign = reg.get('sign_es', reg.get('sign', ''))
            reg_house = ROMAN[reg.get('house', 1) - 1]
            txt = ("La cúspide de tu Casa %s cae en el %s decanato de %s "
                   "(grados %d a %d), cuyo regente es %s. En esta carta, %s está "
                   "en %s, en la Casa %s: observa con qué personas de %s y en los "
                   "temas de la Casa %s puedes trabajar el ámbito de esta casa."
                   % (h.get('roman', ''), _ORD_DECAN[dec], h.get('sign_es', ''),
                      dec * 10, dec * 10 + 10, reg_name, reg_name, reg_sign,
                      reg_house, reg_sign, reg_house))
        else:
            txt = ("La cúspide de tu Casa %s cae en el %s decanato de %s "
                   "(grados %d a %d), cuyo regente es %s."
                   % (h.get('roman', ''), _ORD_DECAN[dec], h.get('sign_es', ''),
                      dec * 10, dec * 10 + 10, reg_name))
        items.append((title, [txt]))
    return items


def build_transit_intro(data):
    """Explicación elaborada de los tránsitos: dónde va el Sol, la Luna y el
    regente del signo solar (signo y casa natal) y qué aspecto le hacen a la
    carta base. Devuelve una lista de párrafos."""
    natal = data.get('natal', {})
    trans = data.get('transit', {})
    cross = data.get('cross_aspects', [])
    tpl = {p['key']: p for p in trans.get('planets', [])}
    npl = {p['key']: p for p in natal.get('planets', [])}
    key_by_name = {v: k for k, v in _NAME_BY_KEY.items()}

    def asp_txt(tk, nk):
        for x in cross:
            if x.get('transit') == tk and x.get('natal') == nk:
                return ", formando %s con tu %s natal (orbe %.1f°)" % (
                    x.get('type_es', 'aspecto').lower(),
                    _NAME_BY_KEY.get(nk, nk), x.get('orb', 0))
        return ""

    def line(tk, label, nk):
        p = tpl.get(tk)
        if not p:
            return ""
        nh = p.get('natal_house', p.get('house', 1))
        return "%s transita por %s, en tu Casa %s%s." % (
            label, p.get('sign_es', ''), ROMAN[nh - 1], asp_txt(tk, nk))

    paras = []
    l = line('aSol', 'El Sol', 'aSol')
    if l:
        paras.append(l)
    l = line('aLuna', 'La Luna', 'aLuna')
    if l:
        paras.append(l)
    sun_sign_en = npl.get('aSol', {}).get('sign', 'Aries')
    ruler_name = _REGENTE.get(sun_sign_en, 'Sol')
    rk = key_by_name.get(ruler_name)
    if rk and rk not in ('aSol', 'aLuna'):
        l = line(rk, 'Tu regente %s' % ruler_name, rk)
        if l:
            paras.append(l)
    return paras


def build_akashic(chart):
    """Registros Akáshicos = estudio de la Casa XII: el signo de su cúspide,
    su regente (dónde está), los planetas dentro de ella y el regente del
    decanato de su cúspide. Devuelve (pre_blocks, sections)."""
    ak = _akashic()
    interp = _interp()
    nat = interp.get('natal', {})
    pm = interp.get('planet_meanings', {})
    by_name = {p.get('name'): p for p in chart['planets']}
    key_by_name = {v: k for k, v in _NAME_BY_KEY.items()}

    houses = chart.get('houses', [])
    h12 = houses[11] if len(houses) >= 12 else {}
    s12_en = h12.get('sign', 'Pisces')
    s12_es = h12.get('sign_es', s12_en)

    # Regente del signo de la Casa XII
    ruler_name = _REGENTE.get(s12_en, 'Neptuno')
    ruler_p = by_name.get(ruler_name)
    # Regente del decanato de la cúspide de la Casa XII
    si = h12.get('sign_index', 1) - 1
    dec = min(2, int(h12.get('deg', 0)) // 10)
    try:
        dk = _DECAN_RULERS[si][dec]
    except Exception:
        dk = None
    dec_p = {p['key']: p for p in chart['planets']}.get(dk) if dk else None
    dec_name = _NAME_BY_KEY.get(dk, dk) if dk else ''
    # Planetas dentro de la Casa XII
    inside = [p for p in chart['planets'] if p.get('house') == 12]

    # Conjuntos en rojo (signos, casas y planetas involucrados con la Casa XII)
    red_signs = {s12_es}
    red_houses = {'Casa XII'}
    red_planets = set()
    if ruler_p:
        red_signs.add(ruler_p.get('sign_es', ''))
        red_houses.add('Casa ' + ROMAN[ruler_p.get('house', 1) - 1])
        red_planets.add(ruler_name)
    for p in inside:
        red_signs.add(p.get('sign_es', ''))
        red_planets.add(p.get('name'))
    if dec_p:
        red_signs.add(dec_p.get('sign_es', ''))
        red_houses.add('Casa ' + ROMAN[dec_p.get('house', 1) - 1])
        red_planets.add(dec_name)

    # ── Pre-bloques: intro + mandamientos + significado general de la Casa XII
    pre = []
    pre.append(("h2", ak.get('intro_titulo', 'Qué son los Registros Akáshicos')))
    for para in ak.get('intro', []):
        pre.append(("p", para))
    if ak.get('metodologia'):
        pre.append(("p", ak['metodologia']))
    pre.append(("h3", ak.get('mand_titulo', 'Los mandamientos zodiacales')))
    if ak.get('mand_intro'):
        pre.append(("p", ak['mand_intro']))
    for sg, casa, frase in ak.get('mand', []):
        hl = (sg in red_signs) or (casa in red_houses)
        pre.append(("mand", '%s y %s' % (sg, casa), frase, hl))
    if ak.get('mand_planetas_intro'):
        pre.append(("p", ak['mand_planetas_intro']))
    for nm, frase in ak.get('mand_planetas', {}).items():
        pre.append(("mand", nm, frase, nm in red_planets))
    pre.append(("h2", ak.get('casa12_titulo', 'El significado de tu Casa XII')))
    for para in ak.get('casa12', []):
        pre.append(("p", para))
    if ak.get('puente_estudio'):
        pre.append(("p", ak['puente_estudio']))

    # ── Secciones: el estudio personalizado ──────────────────────────────
    sections = []

    def _pin(key, sign_en, house):
        out = []
        if pm.get(key):
            out.append(pm[key])
        t = nat.get('planet_in_sign', {}).get(key, {}).get(sign_en)
        if t:
            out.append(t)
        t = nat.get('planet_in_house', {}).get(key, {}).get(str(house))
        if t:
            out.append(t)
        return out

    # 1. Signo en la cúspide de la Casa XII
    items = []
    cusp_txt = nat.get('cusp_in_sign', {}).get('12', {}).get(s12_en)
    if cusp_txt:
        items.append(("Tu Casa XII con %s en la cúspide" % s12_es, [cusp_txt]))
    items.append(("El signo de tu vida pasada",
                  ["Se dice que el signo que ocupa la cúspide de la Casa XII "
                   "indica el signo que tuviste por ascendente en tu vida "
                   "anterior: en tu caso, %s." % s12_es]))
    sections.append(("El signo en la cúspide de tu Casa XII", items))

    # 2. Regente de la Casa XII y dónde está
    if ruler_p:
        rk = ruler_p['key']
        title = "%s, regente de tu Casa XII, en %s y en la Casa %s" % (
            ruler_name, ruler_p.get('sign_es', ruler_p.get('sign', '')),
            ROMAN[ruler_p.get('house', 1) - 1])
        paras = _pin(rk, ruler_p.get('sign'), ruler_p.get('house', 1))
        if paras:
            sections.append(("El regente de tu Casa XII", [(title, paras)]))

    # 3. Planetas dentro de la Casa XII
    if inside:
        items = []
        for p in inside:
            paras = _pin(p['key'], p.get('sign'), 12)
            if paras:
                items.append(("%s en %s, dentro de tu Casa XII"
                              % (p['name'], p.get('sign_es', p.get('sign', ''))), paras))
        if items:
            sections.append(("Los planetas dentro de tu Casa XII", items))
    else:
        sections.append(("Los planetas dentro de tu Casa XII",
            [("Tu Casa XII sin planetas",
              ["No tienes planetas dentro de tu Casa XII. Esto no la deja vacía "
               "de sentido: su historia se lee, sobre todo, a través del signo "
               "de su cúspide, de su regente y del regente del decanato de esa "
               "cúspide, que estudiamos aquí."])]))

    # 4. Regente del decanato de la cúspide
    if dec_p:
        title = "%s, regente del %s decanato de tu cúspide, en %s y en la Casa %s" % (
            dec_name, _ORD_DECAN[dec], dec_p.get('sign_es', dec_p.get('sign', '')),
            ROMAN[dec_p.get('house', 1) - 1])
        paras = _pin(dec_p['key'], dec_p.get('sign'), dec_p.get('house', 1))
        if paras:
            sections.append(("El regente del decanato de tu Casa XII", [(title, paras)]))

    return pre, sections


_CRUZ = {"Aries": "cardinal", "Cancer": "cardinal", "Libra": "cardinal",
         "Capricorn": "cardinal", "Taurus": "fija", "Leo": "fija",
         "Scorpio": "fija", "Aquarius": "fija", "Gemini": "mutable",
         "Virgo": "mutable", "Sagittarius": "mutable", "Pisces": "mutable"}
_ELEM = {"Aries": "fuego", "Leo": "fuego", "Sagittarius": "fuego",
         "Taurus": "tierra", "Virgo": "tierra", "Capricorn": "tierra",
         "Gemini": "aire", "Libra": "aire", "Aquarius": "aire",
         "Cancer": "agua", "Scorpio": "agua", "Pisces": "agua"}


# ── Secciones interpretativas ───────────────────────────────────────────

def build_sections(chart, chart_type='natal', house_key='house', aspects=None):
    """Construye las secciones interpretativas de una carta.
    chart_type: elige la introducción por tipo (natal/transit/solar_return/…).
    house_key: 'house' (casas propias) o 'natal_house' (para tránsitos).
    aspects: lista de aspectos a interpretar; por defecto los de la carta."""
    interp = _interp()
    nat = interp.get('natal', {})
    fb = interp.get('aspect_fallback', {}).get('natal', {})
    bct = interp.get('by_chart_type', {})
    intro = bct.get(chart_type, {}).get('intro', '') or bct.get('natal', {}).get('intro', '')

    # Encabezados por tipo de carta
    PL_HEAD = {'natal': "Los planetas en tu carta natal",
               'transit': "Los planetas en tránsito",
               'solar_return': "Los planetas de tu retorno solar",
               'progressed': "Los planetas de tu carta progresada",
               'combined': "Los planetas de la carta combinada"}
    ASP_HEAD = {'transit': "Aspectos del tránsito a tu carta natal"}

    sections = []
    items = []
    for p in chart['planets']:
        key, sign_en = p['key'], p['sign']
        house = p.get(house_key, p.get('house', 1))
        title = "%s en %s · Casa %s" % (p['name'], p.get('sign_es', sign_en), ROMAN[house - 1])
        paras = []
        t_sign = nat.get('planet_in_sign', {}).get(key, {}).get(sign_en)
        if t_sign:
            paras.append(t_sign)
        t_house = nat.get('planet_in_house', {}).get(key, {}).get(str(house))
        if t_house:
            paras.append(t_house)
        if paras:
            items.append((title, paras))
    if items:
        sections.append((PL_HEAD.get(chart_type, PL_HEAD['natal']), items))

    cusp_items = []
    for h in chart['houses']:
        t = nat.get('cusp_in_sign', {}).get(str(h['num']), {}).get(h['sign'])
        if t:
            cusp_items.append(("Casa %s en %s" % (h['roman'], h.get('sign_es', h['sign'])), [t]))
    if cusp_items:
        sections.append(("Las cúspides de las casas", cusp_items))

    # Decanato de cada cúspide y ubicación de su regente (en todos los reportes)
    dec_items = build_decanates(chart)
    if dec_items:
        sections.append(("Los decanatos de tus casas y sus regentes", dec_items))

    asp = nat.get('aspects', {})
    name_es = {p['key']: p['name'] for p in chart['planets']}
    asp_es = {'Conjunction': 'conjunción', 'Opposition': 'oposición', 'Trine': 'trígono',
              'Square': 'cuadratura', 'Sextile': 'sextil'}
    asp_items = []
    asp_source = aspects if aspects is not None else chart['aspects']
    for a in asp_source:
        ka, kb, typ = a['a'], a['b'], a['type']
        txt = (asp.get(ka, {}).get(typ, {}).get(kb)
               or asp.get(kb, {}).get(typ, {}).get(ka)
               or fb.get(typ))
        if not txt:
            continue
        na = name_es.get(ka) or _NAME_BY_KEY.get(ka, ka)
        nb = name_es.get(kb) or _NAME_BY_KEY.get(kb, kb)
        if chart_type == 'transit':
            title = "%s (tránsito) %s %s (natal)  (orbe %.1f°)" % (
                na, asp_es.get(typ, typ.lower()), nb, a.get('orb', 0))
        else:
            title = "%s %s %s  (orbe %.1f°)" % (
                na, asp_es.get(typ, typ.lower()), nb, a.get('orb', 0))
        asp_items.append((title, [txt]))
    if asp_items:
        sections.append((ASP_HEAD.get(chart_type, "Aspectos astrales"), asp_items))

    return intro, sections


# ── Preámbulo ───────────────────────────────────────────────────────────

def build_preamble(chart, chart_type='natal'):
    pre = _preamble()
    if not pre:
        return []
    planets = {p['key']: p for p in chart['planets']}
    by_name = {p.get('name'): p for p in chart['planets']}
    sun = planets.get('aSol', {})
    moon = planets.get('aLuna', {})
    sun_sign_en = sun.get('sign', 'Aries')
    sun_sign = sun.get('sign_es', sun_sign_en)
    moon_sign = moon.get('sign_es', '')
    asc_sign = chart['angles']['asc'].get('sign_es', '')
    sun_roman = ROMAN[sun.get('house', 1) - 1]
    moon_roman = ROMAN[moon.get('house', 1) - 1]
    sun_house = "Casa " + sun_roman
    moon_house = "Casa " + moon_roman

    # Tierra = signo opuesto al Sol (la Tierra no se calcula como planeta)
    try:
        earth_sign = _SIGNS_ES[(_SIGNS_EN.index(sun_sign_en) + 6) % 12]
    except ValueError:
        earth_sign = ''

    # Regente del signo solar: su nombre, su signo y su casa
    regente = _REGENTE.get(sun_sign_en, 'Sol')
    reg_p = by_name.get(regente, {})
    ruler_sign = reg_p.get('sign_es', '')
    ruler_house = ("Casa " + ROMAN[reg_p.get('house', 1) - 1]) if reg_p else ''

    # Resaltados en rojo (línea completa) — Sol, Luna, Tierra y regente solar
    hl_signos = {s for s in (sun_sign, moon_sign, earth_sign, ruler_sign) if s}
    hl_casas = {h for h in (sun_house, moon_house, ruler_house) if h}
    hl_planetas = {p for p in ("Sol", "Luna", "Tierra", regente) if p}

    b = []
    if chart_type == 'solar_return':
        b.append(("h2", "Tu retorno solar y tu esencia"))
        b.append(("p", "Tu retorno solar conserva el Sol en %s, el mismo signo de tu "
                        "esencia; por eso este ciclo reafirma la cruz y el elemento que "
                        "elegiste, ahora aplicados al año que comienza. Aquí tienes, "
                        "adaptados a este retorno, los mandamientos zodiacales y las "
                        "claves de tu cruz y tu elemento." % sun_sign))
    elif chart_type == 'combined':
        b.append(("h2", "El Sol de la carta combinada y su esencia"))
        b.append(("p", "El Sol de esta carta combinada está en %s. A partir de él se "
                        "definen la cruz y el elemento del vínculo, con sus "
                        "mandamientos. Aquí los tienes, adaptados a la relación que "
                        "ambas personas forman." % sun_sign))
    elif chart_type == 'progressed':
        b.append(("h2", "Tu carta progresada"))
        b.append(("p", "La carta progresada describe tu evolución interna: cómo han "
                        "madurado, paso a paso, las energías con las que naciste. En "
                        "ella, tu Sol progresado está en %s, en la Casa %s; tu Luna "
                        "progresada está en %s, en la Casa %s; y tu Ascendente "
                        "progresado es %s. A partir de estos tres factores se leen la "
                        "cruz y el elemento de esta etapa, con sus mandamientos, que "
                        "encuentras a continuación."
                  % (sun_sign, sun_roman, moon_sign, moon_roman, asc_sign)))
    else:
        b.append(("h2", pre.get('intro_titulo', 'Qué es una carta astral natal')))
        for p in pre.get('intro', []):
            if p.startswith('1.'):
                p += " En tu caso, tu Sol está en %s, en la Casa %s." % (sun_sign, sun_roman)
            elif p.startswith('2.'):
                p += " En tu caso, tu Ascendente es %s." % asc_sign
            elif p.startswith('3.'):
                p += " En tu caso, tu Luna está en %s, en la Casa %s." % (moon_sign, moon_roman)
            b.append(("p", p))
    b.append(("h3", "Los mandamientos zodiacales"))
    b.append(("p", pre.get('mand_signos_intro', '')))
    for k, v in pre.get('mand_signos', {}).items():
        b.append(("mand", k, v, k in hl_signos))
    b.append(("p", pre.get('mand_casas_intro', '')))
    for k, v in pre.get('mand_casas', {}).items():
        b.append(("mand", k, v, k in hl_casas))
    b.append(("p", pre.get('mand_planetas_intro', '')))
    for k, v in pre.get('mand_planetas', {}).items():
        b.append(("mand", k, v, k in hl_planetas))
    b.append(("p", pre.get('mand_cierre', '')))
    b.append(("p", pre.get('puente_cruz', '')))

    cz = pre.get('cruces', {}).get(_CRUZ.get(sun_sign_en, 'fija'), {})
    if cz:
        otros = [s for s in cz.get('signos', []) if s != sun_sign]
        b.append(("h3", cz.get('titulo', 'Tu cruz zodiacal')))
        b.append(("p", "Inicio tu carta por decirte que, al ser %s, perteneces a una "
                        "cruz conocida como %s, la cual decidiste antes de nacer que "
                        "la tienes que formar en esta vida con las personas de los "
                        "signos %s. Primero te digo qué es esta cruz zodiacal y luego "
                        "vamos a tu caso particular."
                  % (sun_sign, cz.get('titulo', '').upper(),
                     ", ".join(otros[:-1]) + " y " + otros[-1] if len(otros) > 1 else "".join(otros))))
        if cz.get('mision'):
            b.append(("p", cz['mision']))
        rep = cz.get('representacion', {})
        if rep:
            b.append(("p", "Representación particular: recuerda que tienes que formar "
                            "esta cruz con personas de los siguientes signos."))
            for sname, stext in rep.items():
                b.append(("p", stext))
        if cz.get('cierre'):
            b.append(("p", cz['cierre']))

    el = pre.get('elementos', {}).get(_ELEM.get(sun_sign_en, 'tierra'), {})
    if el:
        otros2 = [s for s in el.get('signos', []) if s != sun_sign]
        b.append(("h3", el.get('titulo', 'Tu elemento')))
        b.append(("p", "Como pudiste apreciar, %s pertenece al elemento %s de esa "
                        "cruz y, por lo tanto, también decidiste antes de nacer que "
                        "tienes que formar tu triángulo con los signos %s. Observa, "
                        "entonces, qué significa el elemento al cual perteneces."
                  % (sun_sign, el.get('titulo', '').replace('Elemento ', ''),
                     " y ".join(otros2))))
        for fld in ('mision', 'personaje', 'frase', 'descripcion'):
            if el.get(fld):
                b.append(("p", el[fld]))
        for sname, stext in el.get('representacion', {}).items():
            b.append(("p", stext))

    if pre.get('cierre_preambulo'):
        b.append(("p", pre['cierre_preambulo']))
    if pre.get('puente_interpretacion'):
        b.append(("p", pre['puente_interpretacion']))
    return [x for x in b if not (x[0] == "p" and not x[1])]


# ── Glosario ────────────────────────────────────────────────────────────

def _glyph_maps():
    """Mapas nombre→(glifo, color) para signos y planetas, desde report_chart."""
    import unicodedata

    def nk(s):
        return ''.join(c for c in unicodedata.normalize('NFKD', (s or '').lower())
                       if c.isalnum())
    sgn, pln = {}, {}
    if _rc is not None:
        for i, nm in enumerate(_rc.SIGN_NAME):
            sgn[nk(nm)] = (_rc.SIGN_GLYPH[i], _rc.SIGN_COL[i])
        for key, (gl, col, nm) in _rc.PLANETS.items():
            pln[nk(nm)] = (gl, col)
        # variantes de nombre que puede traer el glosario
        pln[nk('R. Fortuna')] = _rc.PLANETS['aRuedaFortuna'][:2]
        pln[nk('Rueda de la Fortuna')] = _rc.PLANETS['aRuedaFortuna'][:2]
        pln[nk('Lilith')] = _rc.PLANETS['aLunaNegra'][:2]
    return sgn, pln


def _glyph_tag(glyph, color):
    # Usa el nombre de fuente REALMENTE registrado (evita romper si falta el TTF
    # de glifos en el servidor). Si no hay fuente de glifos, omite el glifo.
    fn = _rc.GLYPHF() if _rc is not None else None
    if not fn or fn == 'Helvetica':
        return ''
    return ('<font name="%s" color="%s">%s</font>&#160;&#160;' % (fn, color, glyph))


def build_glossary():
    g = _glossary()
    if not g:
        return []
    sgn, pln = _glyph_maps()
    import unicodedata

    def nk(s):
        return ''.join(c for c in unicodedata.normalize('NFKD', (s or '').lower())
                       if c.isalnum())
    b = [("h2", "Glosario")]
    b.append(("h3", "Signos zodiacales"))
    if g.get('signos_intro'):
        b.append(("p", g['signos_intro']))
    for s in g.get('signos', []):
        gc = sgn.get(nk(s['nombre']))
        pre = _glyph_tag(*gc) if gc else ''
        b.append(("item", s['nombre'],
                  "%s · %s. %s" % (s.get('lema', ''), s.get('elemento', ''), s.get('desc', '')),
                  False, pre))
    b.append(("h3", "Luminarias y planetas"))
    if g.get('planetas_intro'):
        b.append(("p", g['planetas_intro']))
    for s in g.get('planetas', []):
        gc = pln.get(nk(s['nombre']))
        pre = _glyph_tag(*gc) if gc else ''
        b.append(("item", s['nombre'], s.get('desc', ''), False, pre))
    b.append(("h3", "Casas astrales"))
    if g.get('casas_intro'):
        b.append(("p", g['casas_intro']))
    for s in g.get('casas', []):
        b.append(("item", s['nombre'], s.get('desc', ''), False))
    return b


_CLOSING_NOUN = {
    'natal': 'Carta Natal', 'solar_return': 'Retorno Solar',
    'progressed': 'Carta Progresada', 'combined': 'Carta Combinada',
    'transit': 'lectura de Tránsitos', 'akashic': 'Registros Akáshicos',
}


def build_closing(chart_type='natal'):
    """Cierre atractivo del reporte, adaptado al tipo de carta."""
    noun = _CLOSING_NOUN.get(chart_type, 'Carta Natal')
    art = 'tus' if chart_type == 'akashic' else 'tu'
    g = _glossary()
    consulta = (g.get('consulta_final')
                or "Esta lectura es un resumen; si deseas profundizar, por favor "
                   "solicita tu consulta en nuestra página web.")
    txt = ("He de dejar aquí el estudio de " + art + " %s, no sin antes recordarte que "
           "TODOS los signos mencionados aquí han de aparecer en algún momento de "
           "tu vida, con personas que nacieron bajo ellos y que en esta existencia "
           "tienen un compromiso vital contigo. Ninguna de ellas es culpable de lo "
           "que te suceda, pero sí responsable de lo que todos escribieron como su "
           "libreto de vida; espero que sepas transformar todo aquello que te "
           "produzca el encuentro con ellas en esta encarnación." % noun)
    return [("h2", "Palabras finales"), ("pc", txt), ("pc", consulta)]


_LEGEND_PLANETS = [("aSol", "Sol"), ("aLuna", "Luna"), ("aMercurio", "Mercurio"),
    ("aVenus", "Venus"), ("aTierra", "Tierra"), ("aMarte", "Marte"),
    ("aJupiter", "Júpiter"), ("aSaturno", "Saturno"),
    ("aChiron", "Quirón"), ("aUrano", "Urano"), ("aNeptuno", "Neptuno"),
    ("aPluton", "Plutón"), ("aNoduloNorte", "Nodo Norte"),
    ("aNoduloSur", "Nodo Sur"), ("aLunaNegra", "Luna Negra"),
    ("aRuedaFortuna", "R. Fortuna")]
_LEGEND_SIGNS = [("Aries", "Aries"), ("Taurus", "Tauro"), ("Gemini", "Géminis"),
    ("Cancer", "Cáncer"), ("Leo", "Leo"), ("Virgo", "Virgo"), ("Libra", "Libra"),
    ("Scorpio", "Escorpio"), ("Sagittarius", "Sagitario"),
    ("Capricorn", "Capricornio"), ("Aquarius", "Acuario"), ("Pisces", "Piscis")]


def build_legend():
    out = []
    for gname, entries in (("Planetas y puntos", _LEGEND_PLANETS),
                           ("Signos", _LEGEND_SIGNS)):
        rows = []
        for key, label in entries:
            p = os.path.join(_ICONS_DIR, key + '.png')
            if os.path.exists(p):
                rows.append((p, label))
        out.append((gname, rows))
    return out


def _png_from_dataurl(chart_png, max_side=900):
    """Decodifica la imagen y la reduce si viene muy grande. Acotar el tamaño
    mantiene bajo el uso de memoria y el tiempo de generación en el servidor."""
    if not chart_png:
        return None
    try:
        if ',' in chart_png:
            chart_png = chart_png.split(',', 1)[1]
        raw = base64.b64decode(chart_png)
    except Exception:
        return None
    try:
        from PIL import Image as PILImage
        im = PILImage.open(io.BytesIO(raw))
        if max(im.size) > max_side:
            ratio = max_side / float(max(im.size))
            im = im.convert('RGB').resize(
                (max(1, int(im.size[0] * ratio)), max(1, int(im.size[1] * ratio))),
                PILImage.LANCZOS)
            b = io.BytesIO()
            im.save(b, 'PNG', optimize=True)
            raw = b.getvalue()
        im.close()
    except Exception:
        pass
    return raw


# ════════════════════════════════════════════════════════════════════════
#  PDF
# ════════════════════════════════════════════════════════════════════════

def render_pdf(sections, chart_png_bytes, meta,
               pre_blocks=None, legend=None, glossary=None, chart_png2_bytes=None,
               closing=None):
    """meta: dict(name, city, astrologer, date, time)"""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors as C
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame,
                                    NextPageTemplate, Paragraph, Spacer,
                                    PageBreak, Image, Table, TableStyle)
    from reportlab.platypus.tableofcontents import TableOfContents

    buf = io.BytesIO()
    page_w, page_h = A4
    lm = rm = 2.6 * cm
    tm = bm = 2.4 * cm
    bf, bd, it = 'Helvetica', 'Helvetica-Bold', 'Helvetica-Oblique'
    # Serif elegante para la portada (siempre disponible en PDF)
    sf, sb, si = 'Times-Roman', 'Times-Bold', 'Times-Italic'
    # Serif elegante para los TÍTULOS del reporte (Cormorant si está; si no, Times)
    if _rc is not None:
        try:
            _rc._reg()   # registra AstroGlyph (glifos) y Cormorant/Times
        except Exception:
            pass
    hser = 'Times-Bold'; hser_i = 'Times-Italic'
    if _rc is not None:
        _s = _rc.SERIF_SB()
        if _s and _s != 'Times-Roman':
            hser = _s
        _si = _rc.SERIF_I()
        if _si:
            hser_i = _si

    person = meta.get('name') or ''
    astrologer = meta.get('astrologer') or _brand().get('astrologo_default', '')
    hoy = fecha_es(_date.today().isoformat())
    cover_img = _cover_path(meta.get('chart_type'))   # portada de marca (o None)

    def esc(t):
        return t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    def with_icons(t, size=12):
        parts = []
        va = -round(size * 0.22)
        for seg, icon in _segments(t):
            parts.append(esc(seg))
            if icon:
                parts.append(' <img src="%s" width="%d" height="%d" valign="%d"/>'
                             % (icon, size, size, va))
        return ''.join(parts)

    # Estilos (capítulos centrados)
    # Tipografía más grande y aireada (lectores adultos mayores)
    st_h2 = ParagraphStyle('h2', fontName=hser, fontSize=21, leading=26,
                           textColor=C.HexColor(NAVY), spaceBefore=20, spaceAfter=12,
                           keepWithNext=1, alignment=TA_CENTER)
    st_h3 = ParagraphStyle('h3', fontName=hser, fontSize=16, leading=20,
                           textColor=C.HexColor(BLUE), spaceBefore=12, spaceAfter=6,
                           keepWithNext=1)
    st_body = ParagraphStyle('b', fontName=bf, fontSize=12, leading=18.5,
                             textColor=C.HexColor(INK), alignment=TA_JUSTIFY,
                             spaceAfter=12, firstLineIndent=16)
    st_item = ParagraphStyle('li', fontName=bf, fontSize=12, leading=18.5,
                             textColor=C.HexColor(INK), leftIndent=14, spaceAfter=3)

    _WM_SIGNS = ['Aries', 'Taurus', 'Gemini', 'Cancer', 'Leo', 'Virgo', 'Libra',
                 'Scorpio', 'Sagittarius', 'Capricorn', 'Aquarius', 'Pisces']

    def _watermark_signs(cv):
        """Los 12 signos, sutiles y dorados, repartidos alrededor del margen."""
        import math
        sz = 24
        # Los signos van holgadamente por dentro del marco dorado (que está a
        # 44 pt del borde) y por encima del pie de portada, para que ninguno
        # quede sobre una línea ni sobre un texto.
        inset = 44 + sz / 2 + 12
        top_y = page_h - inset          # justo dentro del marco superior
        bot_y = 145                     # por encima de astrólogo/fecha/marca
        cx = page_w / 2
        cy = (top_y + bot_y) / 2.0
        rx = page_w / 2 - inset
        ry = (top_y - bot_y) / 2.0
        for i, sg in enumerate(_WM_SIGNS):
            fp = os.path.join(_ICONS_DIR, 'wm_' + sg + '.png')
            if not os.path.exists(fp):
                continue
            ang = math.pi / 2 - i * (2 * math.pi / 12)  # arranca arriba, en sentido horario
            x = cx + rx * math.cos(ang)
            y = cy + ry * math.sin(ang)
            try:
                cv.drawImage(fp, x - sz / 2, y - sz / 2, width=sz, height=sz, mask='auto')
            except Exception:
                pass

    def on_cover(cv, doc):
        # Si hay portada de marca para este tipo, va a página completa
        if cover_img:
            try:
                # Fondo oscuro + imagen a altura completa sin deformar (la
                # portada es vertical 9:16; se centra con finas bandas laterales)
                cv.setFillColor(C.HexColor('#0A0A16'))
                cv.rect(0, 0, page_w, page_h, fill=1, stroke=0)
                cv.drawImage(cover_img, 0, 0, width=page_w, height=page_h,
                             preserveAspectRatio=True, anchor='c', mask='auto')
                return
            except Exception:
                pass
        cv.saveState()
        cv.setFillColorRGB(0x0E / 255., 0x13 / 255., 0x32 / 255.)
        cv.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        _watermark_signs(cv)
        cv.setStrokeColor(C.HexColor(GOLD))
        cv.setLineWidth(1.0)
        cv.rect(36, 36, page_w - 72, page_h - 72, fill=0, stroke=1)
        cv.setLineWidth(0.4)
        cv.rect(44, 44, page_w - 88, page_h - 88, fill=0, stroke=1)
        # ornamento
        cx, cy, r = page_w / 2, page_h * 0.76, 9
        p = cv.beginPath()
        p.moveTo(cx, cy + r); p.lineTo(cx + r, cy); p.lineTo(cx, cy - r)
        p.lineTo(cx - r, cy); p.close()
        cv.setLineWidth(0.9)
        cv.drawPath(p, fill=0, stroke=1)
        cv.circle(cx, cy, 1.6, fill=1, stroke=0)
        cv.setLineWidth(0.5)
        cv.line(cx - 110, cy, cx - r - 12, cy)
        cv.line(cx + r + 12, cy, cx + 110, cy)
        # astrólogo y fecha de generación
        if astrologer:
            cv.setFont(si, 11.5)
            cv.setFillColor(C.HexColor(LIGHT))
            cv.drawCentredString(page_w / 2, 118, 'Astrólogo: %s' % astrologer)
        cv.setFont(sf, 9.5)
        cv.setFillColor(C.HexColor(SLATE))
        cv.drawCentredString(page_w / 2, 100, hoy)
        cv.setFont(bf, 9)
        cv.setFillColor(C.HexColor(GOLD))
        cv.drawCentredString(page_w / 2, 58, 'R E P O R T E   A S T R A L')
        cv.restoreState()

    def _diamond(cv, cx, cy, r):
        p = cv.beginPath()
        p.moveTo(cx, cy + r); p.lineTo(cx + r, cy); p.lineTo(cx, cy - r)
        p.lineTo(cx - r, cy); p.close()
        cv.drawPath(p, fill=0, stroke=1)

    _redes = _brand().get('redes', [])
    _web = next((r.get('texto') for r in _redes if r.get('icon') == 'web'),
                'www.mauriciopuerta.tv')

    def _footer_sub():
        _tipo = meta.get('subtitle', 'de Carta Natal')
        if _tipo.lower().startswith('de '):
            _tipo = _tipo[3:]
        parts = [_tipo]
        if astrologer:
            parts.append(astrologer)
        if _web:
            parts.append(_web)
        return '   ·   ' + '   ·   '.join(parts)

    def _header_line():
        """Encabezado con los datos del consultante (todas las páginas)."""
        parts = ['REPORTE ASTRAL']
        if person:
            parts.append(person)
        dt = []
        if meta.get('time'):
            dt.append(meta['time'])
        fa = fecha_ab(meta.get('date', ''))
        if fa:
            dt.append(fa)
        if meta.get('city'):
            dt.append(meta['city'])
        if dt:
            parts.append(', '.join(dt))
        return '   —   '.join(parts)

    def _draw_header(cv, wpage):
        # Gris tenue (marca de agua); nombre, hora y fecha en cursiva
        gray = C.HexColor('#AEB4C6')
        left = 'REPORTE ASTRAL'
        datos = []
        if person:
            datos.append(person)
        dt = []
        if meta.get('time'):
            dt.append(meta['time'])
        fa = fecha_ab(meta.get('date', ''))
        if fa:
            dt.append(fa)
        if meta.get('city'):
            dt.append(meta['city'])
        if dt:
            datos.append(', '.join(dt))
        right = ('   —   ' + '   —   '.join(datos)) if datos else ''
        w1 = cv.stringWidth(left, bf, 8)
        w2 = cv.stringWidth(right, it, 8)
        x0 = (wpage - (w1 + w2)) / 2.0
        y = page_h - tm * 0.32
        cv.setFillColor(gray)
        cv.setFont(bf, 8); cv.drawString(x0, y, left)
        cv.setFont(it, 8); cv.drawString(x0 + w1, y, right)

    def _draw_footer(cv, wpage, doc):
        """Pie: línea, marca · tipo · web y número de página (en rombo)."""
        cv.saveState()
        cv.setStrokeColor(C.HexColor(GOLD)); cv.setLineWidth(0.7)
        cv.line(lm, bm * 0.72, wpage - rm, bm * 0.72)
        marca = 'REPORTE ASTRAL'
        cv.setFont(bf, 7.5); cv.setFillColor(C.HexColor(NAVY))
        cv.drawString(lm, bm * 0.42, marca)
        cv.setFont(it, 7.5); cv.setFillColor(C.HexColor(SLATE))
        cv.drawString(lm + cv.stringWidth(marca, bf, 7.5), bm * 0.42, _footer_sub())
        cx2 = wpage - rm - 8; cy2 = bm * 0.46
        cv.setStrokeColor(C.HexColor(GOLD)); cv.setLineWidth(0.7)
        _diamond(cv, cx2, cy2, 8.5)
        cv.setFont(bf, 7.5); cv.setFillColor(C.HexColor(NAVY))
        cv.drawCentredString(cx2, cy2 - 2.6, str(doc.page))
        cv.restoreState()

    def on_body(cv, doc):
        cv.saveState()
        cv.setFillColor(C.HexColor(PAPER_BG)); cv.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        _draw_header(cv, page_w)
        ty = page_h - tm * 0.62
        # Divisor al estilo de las nuevas cartas: línea dorada + rombo relleno
        cx_ = page_w / 2
        cv.setStrokeColor(C.HexColor(GOLD)); cv.setLineWidth(0.7)
        cv.line(cx_ - 94, ty, cx_ - 8, ty)
        cv.line(cx_ + 8, ty, cx_ + 94, ty)
        pd = cv.beginPath()
        pd.moveTo(cx_, ty + 3.4); pd.lineTo(cx_ + 3.4, ty)
        pd.lineTo(cx_, ty - 3.4); pd.lineTo(cx_ - 3.4, ty); pd.close()
        cv.setFillColor(C.HexColor(GOLD)); cv.drawPath(pd, fill=1, stroke=0)
        cv.setLineWidth(0.7)
        cv.line(lm, bm * 0.72, page_w - rm, bm * 0.72)
        cv.setFont(bf, 7.5); cv.setFillColor(C.HexColor(NAVY))
        marca = 'REPORTE ASTRAL'
        cv.drawString(lm, bm * 0.42, marca)
        cv.setFont(it, 7.5); cv.setFillColor(C.HexColor(SLATE))
        sub = _footer_sub()
        cv.drawString(lm + cv.stringWidth(marca, bf, 7.5), bm * 0.42, sub)
        num = str(doc.page)
        cx2 = page_w - rm - 8
        cy2 = bm * 0.46
        cv.setStrokeColor(C.HexColor(GOLD)); cv.setLineWidth(0.7)
        _diamond(cv, cx2, cy2, 8.5)
        cv.setFont(bf, 7.5); cv.setFillColor(C.HexColor(NAVY))
        cv.drawCentredString(cx2, cy2 - 2.6, num)
        cv.restoreState()

    # Página apaisada (para comparar las dos ruedas lado a lado, grandes)
    land_w, land_h = landscape(A4)

    def on_wheels(cv, doc):
        cv.saveState()
        cv.setFillColor(C.HexColor(PAPER_BG)); cv.rect(0, 0, land_w, land_h, fill=1, stroke=0)
        cv.setFont(bf, 8); cv.setFillColor(C.HexColor(NAVY))
        cv.drawCentredString(land_w / 2, land_h - tm * 0.30, _header_line())
        cv.setStrokeColor(C.HexColor(GOLD)); cv.setLineWidth(0.7)
        cv.line(lm, bm * 0.72, land_w - rm, bm * 0.72)
        cv.setFont(bf, 7.5); cv.setFillColor(C.HexColor(NAVY))
        marca = 'REPORTE ASTRAL'
        cv.drawString(lm, bm * 0.42, marca)
        cv.setFont(it, 7.5); cv.setFillColor(C.HexColor(SLATE))
        cv.drawString(lm + cv.stringWidth(marca, bf, 7.5), bm * 0.42, _footer_sub())
        cx2 = land_w - rm - 8
        cv.setStrokeColor(C.HexColor(GOLD)); cv.setLineWidth(0.7)
        _diamond(cv, cx2, bm * 0.46, 8.5)
        cv.setFont(bf, 7.5); cv.setFillColor(C.HexColor(NAVY))
        cv.drawCentredString(cx2, bm * 0.46 - 2.6, str(doc.page))
        cv.restoreState()

    class _Doc(BaseDocTemplate):
        def afterFlowable(self, f):
            key = getattr(f, '_tocKey', None)
            if key:
                lvl = getattr(f, '_tocLevel', None)
                if lvl is None:
                    lvl = 0 if getattr(getattr(f, 'style', None), 'name', '') == 'h2' else 1
                txt = getattr(f, '_tocText', None)
                if not txt:
                    txt = f.getPlainText() if hasattr(f, 'getPlainText') else ''
                self.notify('TOCEntry', (lvl, txt, self.page, key))

    frame = Frame(lm, bm, page_w - lm - rm, page_h - tm - bm)
    frame_land = Frame(lm, bm, land_w - lm - rm, land_h - tm * 0.7 - bm)
    frame_full = Frame(0, 0, page_w, page_h, leftPadding=0, rightPadding=0,
                       topPadding=0, bottomPadding=0)

    def on_plate(cv, doc):
        cv.saveState()
        cv.setFillColor(C.HexColor('#FFFFFF'))
        cv.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        cv.restoreState()

    # ── Lámina estética de carta (nueva) + página de leyenda ──────────────
    use_lamina = bool(_rc is not None and meta.get('lamina_chart') is not None)

    def on_lamina_chart(cv, doc):
        try:
            _rc.draw_lamina(cv, page_w, page_h, meta.get('chart_type', 'natal'),
                            meta, meta['lamina_chart'])
        except Exception:
            cv.saveState(); cv.setFillColor(C.HexColor('#FFFFFF'))
            cv.rect(0, 0, page_w, page_h, fill=1, stroke=0); cv.restoreState()
        if meta.get('_key_chart'):
            cv.bookmarkPage(meta['_key_chart'])
        _draw_footer(cv, page_w, doc)

    def on_lamina_legend(cv, doc):
        try:
            _rc.draw_legend(cv, page_w, page_h)
        except Exception:
            cv.saveState(); cv.setFillColor(C.HexColor('#FFFFFF'))
            cv.rect(0, 0, page_w, page_h, fill=1, stroke=0); cv.restoreState()
        if meta.get('_key_legend'):
            cv.bookmarkPage(meta['_key_legend'])
        _draw_footer(cv, page_w, doc)

    def on_lamina_dual(cv, doc):
        try:
            caps = meta.get('dual_captions') or ['Carta natal', '']
            _rc.draw_dual_wheels(cv, land_w, land_h, meta,
                                 meta.get('lamina_natal'), meta.get('lamina_chart'),
                                 caps[0], caps[1] if len(caps) > 1 else '')
        except Exception:
            cv.saveState(); cv.setFillColor(C.HexColor('#FFFFFF'))
            cv.rect(0, 0, land_w, land_h, fill=1, stroke=0); cv.restoreState()
        if meta.get('_key_dual'):
            cv.bookmarkPage(meta['_key_dual'])
        _draw_footer(cv, land_w, doc)

    doc = _Doc(buf, pagesize=A4, leftMargin=lm, rightMargin=rm,
               topMargin=tm, bottomMargin=bm)
    doc.addPageTemplates([
        PageTemplate(id='Cover', frames=[Frame(lm, bm, page_w - lm - rm, page_h - tm - bm)], onPage=on_cover),
        PageTemplate(id='Body', frames=[frame], onPage=on_body),
        PageTemplate(id='Wheels', frames=[frame_land], onPage=on_wheels, pagesize=landscape(A4)),
        PageTemplate(id='Plate', frames=[frame_full], onPage=on_plate),
        PageTemplate(id='LaminaChart', frames=[Frame(0, 0, page_w, page_h, leftPadding=0,
                     rightPadding=0, topPadding=0, bottomPadding=0)], onPage=on_lamina_chart),
        PageTemplate(id='LaminaLegend', frames=[Frame(0, 0, page_w, page_h, leftPadding=0,
                     rightPadding=0, topPadding=0, bottomPadding=0)], onPage=on_lamina_legend),
        PageTemplate(id='LaminaDual', frames=[Frame(0, 0, land_w, land_h, leftPadding=0,
                     rightPadding=0, topPadding=0, bottomPadding=0)], onPage=on_lamina_dual,
                     pagesize=landscape(A4)),
    ])

    nkey = [0]
    content = []

    def H2(text):
        key = "sec%d" % nkey[0]; nkey[0] += 1
        p = Paragraph('<a name="%s"/>' % key + with_icons(text, 15), st_h2)
        p._tocKey = key
        content.append(p)

    def H3(text):
        key = "sec%d" % nkey[0]; nkey[0] += 1
        p = Paragraph('<a name="%s"/>' % key + with_icons(text, 14), st_h3)
        p._tocKey = key
        content.append(p)

    # Estilos para los mandamientos tabulados
    st_mand_l = ParagraphStyle('mndl', fontName=bd, fontSize=12, leading=17,
                               textColor=C.HexColor(INK))
    st_mand_r = ParagraphStyle('mndr', fontName=bf, fontSize=12, leading=17,
                               textColor=C.HexColor(INK))
    st_mand_lR = ParagraphStyle('mndlR', fontName=bd, fontSize=12, leading=17,
                                textColor=C.HexColor(RED))
    st_mand_rR = ParagraphStyle('mndrR', fontName=bf, fontSize=12, leading=17,
                                textColor=C.HexColor(RED))

    def flush_mand(rows):
        """Vuelca los mandamientos acumulados como tabla alineada (signo | frase).
        Las líneas resaltadas van completas en rojo."""
        if not rows:
            return
        data_rows = []
        for label, frase, hl in rows:
            lp = Paragraph(with_icons(label, 13), st_mand_lR if hl else st_mand_l)
            rp = Paragraph(esc(frase), st_mand_rR if hl else st_mand_r)
            data_rows.append([lp, rp])
        col0 = 4.6 * cm
        t = Table(data_rows, colWidths=[col0, (page_w - lm - rm) - col0])
        t.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING', (0, 0), (-1, -1), 2.5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
            ('LEFTPADDING', (0, 0), (0, -1), 6),
        ]))
        content.append(t)

    legend_done = [False]

    def add_legend(ncols=4, total_w=None, fs=9, ic_sz=10):
        """Leyenda de símbolos. En la página apaisada se usa con más columnas."""
        if not legend:
            return
        tw = total_w if total_w else (page_w - lm - rm)
        H3("Leyenda de símbolos")
        st_leg = ParagraphStyle('lg%d' % ncols, fontName=bf, fontSize=fs,
                                textColor=C.HexColor(INK))
        for gname, rows in legend:
            content.append(Paragraph('<b>%s</b>' % esc(gname), ParagraphStyle(
                'lgt%d' % ncols, fontName=bd, fontSize=fs + 1,
                textColor=C.HexColor(BLUE), spaceBefore=3, spaceAfter=1)))
            cells, row = [], []
            for icon, label in rows:
                row.append(Paragraph(
                    '<img src="%s" width="%d" height="%d" valign="-2"/> %s'
                    % (icon, ic_sz, ic_sz, esc(label)), st_leg))
                if len(row) == ncols:
                    cells.append(row); row = []
            if row:
                row += [Paragraph('', st_leg)] * (ncols - len(row))
                cells.append(row)
            t = Table(cells, colWidths=[tw / float(ncols)] * ncols)
            t.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            content.append(t)

    if use_lamina:
        # Lámina estética (rueda + datos), luego —para retorno/progresada— una
        # página apaisada con las dos ruedas, y por último la leyenda. Cada una
        # se registra en el índice mediante un Spacer marcador + bookmark.
        meta['_key_chart'] = 'lam_chart'
        meta['_key_dual'] = 'lam_dual'
        meta['_key_legend'] = 'lam_legend'

        sp = Spacer(1, 2)
        sp._tocKey = meta['_key_chart']
        sp._tocText = meta.get('lamina_toc') or meta.get('chart_heading', 'Carta natal')
        sp._tocLevel = 0
        content.append(sp)                                 # ocupa la página LaminaChart
        if meta.get('lamina_natal') is not None:
            content.append(NextPageTemplate('LaminaDual'))
            content.append(PageBreak())
            sp2 = Spacer(1, 2)
            sp2._tocKey = meta['_key_dual']
            sp2._tocText = meta.get('dual_title', 'Las dos cartas')
            sp2._tocLevel = 1
            content.append(sp2)                            # ocupa la página doble apaisada
        content.append(NextPageTemplate('LaminaLegend'))
        content.append(PageBreak())
        sp3 = Spacer(1, 2)
        sp3._tocKey = meta['_key_legend']
        sp3._tocText = 'Leyenda de símbolos'
        sp3._tocLevel = 0
        content.append(sp3)                                # ocupa la página de leyenda
        content.append(NextPageTemplate('Body'))
        content.append(PageBreak())
        legend_done[0] = True
    elif chart_png_bytes:
        try:
            from PIL import Image as PILImage
            caps = meta.get('img_captions')
            if chart_png2_bytes and caps:
                # Ambas ruedas en UNA página apaisada (la orientación la fija
                # el flujo, justo tras el índice). Título compacto (ahorra espacio para que la leyenda entre)
                st_wh = ParagraphStyle('wh', fontName=bd, fontSize=14, leading=17,
                                       textColor=C.HexColor(NAVY), alignment=TA_CENTER,
                                       spaceBefore=0, spaceAfter=4)
                content.append(Paragraph(esc(meta.get('chart_heading', 'Tus dos cartas')), st_wh))
                colw = (land_w - lm - rm) / 2.0
                # Deja sitio para la leyenda en la misma página
                each = min(colw - 0.5 * cm, 8.9 * cm)
                st_wcap = ParagraphStyle('wcap', fontName=bd, fontSize=12, leading=15,
                                         textColor=C.HexColor(BLUE), alignment=TA_CENTER,
                                         spaceAfter=5)
                row = []
                for img_bytes, cap in ((chart_png_bytes, caps[0]),
                                       (chart_png2_bytes, caps[1])):
                    im = PILImage.open(io.BytesIO(img_bytes))
                    iw, ih = im.size
                    w = each; h = w * ih / iw
                    row.append([Paragraph(cap, st_wcap),
                                Image(io.BytesIO(img_bytes), width=w, height=h)])
                tw = Table([row], colWidths=[colw, colw])
                tw.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                        ('ALIGN', (0, 0), (-1, -1), 'CENTER')]))
                content.append(tw)
                # La leyenda entra en la misma página apaisada
                add_legend(ncols=8, total_w=land_w - lm - rm, fs=8, ic_sz=9)
                legend_done[0] = True
                content.append(NextPageTemplate('Body'))
                content.append(PageBreak())
            else:
                im = PILImage.open(io.BytesIO(chart_png_bytes))
                iw, ih = im.size
                # Página dedicada (solo carta + leyenda): aprovechar el máximo
                disp_w = page_w - lm - rm
                disp_h = disp_w * ih / iw
                max_h = 16.4 * cm       # deja sitio seguro para la leyenda debajo
                if disp_h > max_h:
                    disp_h = max_h; disp_w = disp_h * iw / ih
                H2(meta.get('chart_heading', 'Carta natal: el libreto de tu vida'))
                content.append(Spacer(1, 4))
                img_tbl = Table([[Image(io.BytesIO(chart_png_bytes), width=disp_w, height=disp_h)]],
                                colWidths=[page_w - lm - rm])
                img_tbl.setStyle(TableStyle([('ALIGN', (0, 0), (-1, -1), 'CENTER')]))
                content.append(img_tbl)
                content.append(Spacer(1, 4))
        except Exception:
            pass

    if legend and not legend_done[0]:
        add_legend(ncols=6, fs=8.5, ic_sz=9)
        content.append(PageBreak())   # la carta + leyenda quedan solas en su página

    if pre_blocks:
        mand_buf = []
        for blk in pre_blocks:
            if blk[0] == "mand":
                mand_buf.append((blk[1], blk[2], blk[3]))
                continue
            if mand_buf:
                flush_mand(mand_buf); mand_buf = []
                content.append(Spacer(1, 8))
            if blk[0] == "h2":
                H2(blk[1])
            elif blk[0] == "h3":
                H3(blk[1])
            elif blk[0] == "p":
                content.append(Paragraph(with_icons(blk[1], 12), st_body))
            elif blk[0] == "item":
                _, nm, tx, hl = blk
                nm_markup = with_icons(nm, 12)
                nm_markup = ('<font color="%s"><b>%s</b></font>' % (RED, nm_markup)) if hl \
                    else ('<b>%s</b>' % nm_markup)
                content.append(Paragraph(nm_markup + ' — ' + esc(tx), st_item))
        if mand_buf:
            flush_mand(mand_buf); mand_buf = []

    # Título de interpretación fusionado con el inicio de su texto: así el
    # título nunca queda solo al pie y el párrafo se parte con naturalidad,
    # sin dejar huecos grandes.
    st_ihead = ParagraphStyle('ihead', fontName=bf, fontSize=12, leading=18.5,
                              textColor=C.HexColor(INK), alignment=TA_JUSTIFY,
                              spaceBefore=13, spaceAfter=6, firstLineIndent=0)

    def emit_item(title, paras):
        key = "sec%d" % nkey[0]; nkey[0] += 1
        head = ('<a name="%s"/><font size="13" color="%s"><b>%s</b></font>'
                % (key, BLUE, with_icons(title, 14)))
        first = with_icons(paras[0], 12) if paras else ''
        p = Paragraph(head + '<br/><br/>' + first, st_ihead)
        p._tocKey = key; p._tocText = title; p._tocLevel = 1
        content.append(p)
        for extra in paras[1:]:
            content.append(Paragraph(with_icons(extra, 12), st_body))

    for sec_title, items in sections:
        H2(sec_title)
        for it_title, paras in items:
            emit_item(it_title, paras)

    if glossary:
        content.append(Spacer(1, 20))
        for blk in glossary:
            if blk[0] == "h2":
                H2(blk[1])
            elif blk[0] == "h3":
                H3(blk[1])
            elif blk[0] == "p":
                content.append(Paragraph(with_icons(blk[1], 12), st_body))
            elif blk[0] == "item":
                nm = blk[1]; tx = blk[2]
                pre = blk[4] if len(blk) > 4 else ''
                nm_m = esc(nm) if pre else with_icons(nm, 12)
                content.append(Paragraph(pre + '<b>%s</b> — %s' % (nm_m, esc(tx)),
                                         st_item))

    # ── Cierre / palabras finales ────────────────────────────────────────
    if closing:
        st_close = ParagraphStyle('close', fontName=si, fontSize=13, leading=21,
                                  textColor=C.HexColor(NAVY), alignment=TA_CENTER,
                                  spaceBefore=8, spaceAfter=14,
                                  leftIndent=18, rightIndent=18)
        content.append(PageBreak())
        for blk in closing:
            if blk[0] == "h2":
                H2(blk[1])
            elif blk[0] == "pc":
                content.append(Paragraph(esc(blk[1]), st_close))

    # ── Antepenúltima: página de información (rueda educativa cuadrada) ───
    info_img = _info_path()
    if info_img:
        try:
            from PIL import Image as PILImage2
            iw, ih = PILImage2.open(info_img).size
            side = 1.35 * cm
            w = page_w - 2 * side; h = w * ih / iw
            if h > page_h - 2 * side:
                h = page_h - 2 * side; w = h * iw / ih
            content.append(NextPageTemplate('Plate'))
            content.append(PageBreak())
            # Centrado vertical (la imagen cuadrada deja aire arriba y abajo)
            top = max(0, (page_h - h) / 2.0 - 6)
            content.append(Spacer(1, top))
            itbl = Table([[Image(info_img, width=w, height=h)]], colWidths=[page_w])
            itbl.setStyle(TableStyle([('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                      ('LEFTPADDING', (0, 0), (-1, -1), 0),
                                      ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                                      ('TOPPADDING', (0, 0), (-1, -1), 0),
                                      ('BOTTOMPADDING', (0, 0), (-1, -1), 0)]))
            content.append(itbl)
            content.append(NextPageTemplate('Body'))
        except Exception:
            pass

    # ── Penúltima: aviso legal y de privacidad ───────────────────────────
    content.append(PageBreak())
    st_dl_h = ParagraphStyle('dlh', fontName=bd, fontSize=14.5, leading=19,
                             textColor=C.HexColor(NAVY), spaceBefore=18, spaceAfter=6)
    st_dl_p = ParagraphStyle('dlp', fontName=bf, fontSize=12.5, leading=20,
                             textColor=C.HexColor(INK), alignment=TA_JUSTIFY, spaceAfter=16)
    H2("Aviso legal y de privacidad")
    content.append(Spacer(1, 10))
    for htitle, ptext in _DISCLAIMER:
        content.append(Paragraph(esc(htitle), st_dl_h))
        content.append(Paragraph(esc(ptext), st_dl_p))

    # ── Redes sociales del astrólogo (última página) ─────────────────────
    brand = _brand()
    redes = brand.get('redes', [])
    if redes:
        content.append(PageBreak())
        H2(brand.get('redes_titulo', 'Encuéntrame en'))
        st_red = ParagraphStyle('rd', fontName=bf, fontSize=11, leading=22,
                                textColor=C.HexColor(INK), alignment=TA_CENTER)
        if astrologer:
            content.append(Paragraph('<font name="%s" size="13" color="%s"><i>%s</i></font>'
                                     % (si, BLUE, esc(astrologer)), ParagraphStyle(
                                         'rda', alignment=TA_CENTER, spaceAfter=12,
                                         fontName=si, fontSize=13,
                                         textColor=C.HexColor(BLUE))))
        for r in redes:
            ic = os.path.join(_ICONS_DIR, 'soc_' + r.get('icon', '') + '.png')
            ic_markup = ('<img src="%s" width="13" height="13" valign="-2"/> ' % ic) \
                if os.path.exists(ic) else ''
            content.append(Paragraph(
                ic_markup + '<b>%s</b> — <a href="%s" color="%s"><u>%s</u></a>'
                % (esc(r.get('nombre', '')), r.get('url', '#'), BLUE,
                   esc(r.get('texto', r.get('url', '')))), st_red))

        # ── Nuestra app: Diario Astral ───────────────────────────────────
        _app = os.path.join(_COVERS_DIR, 'diario_astral.png')
        content.append(Spacer(1, 26))
        if os.path.exists(_app):
            ap = Table([[Image(_app, width=1.7 * cm, height=1.7 * cm)]],
                       colWidths=[page_w - lm - rm])
            ap.setStyle(TableStyle([('ALIGN', (0, 0), (-1, -1), 'CENTER')]))
            content.append(ap)
        content.append(Paragraph(
            '<b>Diario Astral</b>: tu brújula cósmica personal, en español.',
            ParagraphStyle('appt', fontName=sb, fontSize=13, leading=17,
                           textColor=C.HexColor(NAVY), alignment=TA_CENTER,
                           spaceBefore=6, spaceAfter=3)))
        content.append(Paragraph(
            'Descárgala en App Store y Google Play.',
            ParagraphStyle('apps', fontName=si, fontSize=11, leading=15,
                           textColor=C.HexColor(SLATE), alignment=TA_CENTER)))

    # ── Índice con números de página ─────────────────────────────────────
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle('tocc', fontName=bd, fontSize=11.5, leading=18,
                       textColor=C.HexColor(NAVY)),
        ParagraphStyle('tocs', fontName=bf, fontSize=9.5, leading=14.5,
                       textColor=C.HexColor(BLUE), leftIndent=20),
    ]
    toc.dotsMinLevel = 0

    # ── Portada / página de datos ────────────────────────────────────────
    def build_data_flow(on_dark):
        """Datos de la persona. on_dark=True → sobre la portada diseñada
        (fondo oscuro); on_dark=False → página de datos sobre fondo blanco."""
        c_gold = GOLD
        c_sub = SLATE
        c_name = LIGHT if on_dark else NAVY
        c_line = SLATE if on_dark else INK
        c_ret = GOLD
        s1 = ParagraphStyle('d1', fontName=sb, fontSize=30, leading=36,
                            textColor=C.HexColor(c_gold), alignment=TA_CENTER)
        s2 = ParagraphStyle('d2', fontName=si, fontSize=17, leading=22,
                            textColor=C.HexColor(c_sub), alignment=TA_CENTER)
        s3 = ParagraphStyle('d3', fontName=sb, fontSize=19, leading=25,
                            textColor=C.HexColor(c_name), alignment=TA_CENTER)
        s4 = ParagraphStyle('d4', fontName=sf, fontSize=12.5, leading=19,
                            textColor=C.HexColor(c_line), alignment=TA_CENTER)
        s5 = ParagraphStyle('d5', fontName=si, fontSize=11, leading=16,
                            textColor=C.HexColor(c_ret), alignment=TA_CENTER)
        s6 = ParagraphStyle('d6', fontName=sf, fontSize=11.5, leading=17,
                            textColor=C.HexColor(c_sub), alignment=TA_CENTER)
        f = [Paragraph('Reporte Astrológico', s1), Spacer(1, 4),
             Paragraph(esc(meta.get('subtitle', 'de Carta Natal')), s2), Spacer(1, 34)]
        if person:
            f.append(Paragraph(esc(person), s3)); f.append(Spacer(1, 10))
        lf = meta.get('time', '')
        if lf:
            lf += '  —  '
        lf += fecha_es(meta.get('date', ''))
        f.append(Paragraph(esc(lf), s4))
        if meta.get('city'):
            f.append(Paragraph(esc(meta['city']), s4))
        if meta.get('chart_type') == 'solar_return' and meta.get('return_moment'):
            f.append(Spacer(1, 16))
            if meta.get('relocated') and meta.get('return_place'):
                f.append(Paragraph('Relocalizado en: ' + esc(meta['return_place']), s5))
            rm = esc(meta['return_moment'])
            if meta.get('return_tz'):
                rm += '  (' + esc(meta['return_tz']) + ')'
            f.append(Paragraph('El retorno solar exacto ocurre el', s6))
            f.append(Paragraph(rm, s5))
            if meta.get('return_ut'):
                f.append(Paragraph(esc(meta['return_ut']) + ' UT', s6))
        pb2 = meta.get('person_b')
        if meta.get('chart_type') == 'combined' and pb2:
            f.append(Spacer(1, 14)); f.append(Paragraph('en vínculo con', s6))
            if pb2.get('name'):
                f.append(Paragraph(esc(pb2['name']), s3))
            lb = pb2.get('time', '')
            if lb:
                lb += '  —  '
            lb += fecha_es(pb2.get('date', ''))
            f.append(Paragraph(esc(lb), s4))
            if pb2.get('city'):
                f.append(Paragraph(esc(pb2['city']), s4))
        return f

    if cover_img:
        # Página 1 = portada de marca (imagen); página 2 = datos sobre blanco
        flow = [NextPageTemplate('Body'), PageBreak(), Spacer(1, 70)]
        flow += build_data_flow(on_dark=False)
        flow.append(PageBreak())
    else:
        # Portada diseñada (fondo oscuro) con los datos encima
        flow = [NextPageTemplate('Body'), Spacer(1, page_h * 0.24)]
        flow += build_data_flow(on_dark=True)
        flow.append(PageBreak())

    flow.append(Paragraph('Índice de contenidos', st_h2))
    flow.append(toc)
    # La carta va JUSTO después del índice.
    if use_lamina:
        # Lámina estética a página completa (portrait), no apaisada.
        flow.append(NextPageTemplate('LaminaChart'))
    elif chart_png2_bytes and meta.get('img_captions'):
        # (modo antiguo) dos ruedas → primera página apaisada.
        flow.append(NextPageTemplate('Wheels'))
    flow.append(PageBreak())
    flow += content
    doc.multiBuild(flow)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════════════
#  DOCX
# ════════════════════════════════════════════════════════════════════════

def render_docx(sections, chart_png_bytes, meta,
                pre_blocks=None, legend=None, glossary=None, chart_png2_bytes=None,
                closing=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY_RGB = RGBColor(0x1E, 0x24, 0x50)
    BLUE_RGB = RGBColor(0x3A, 0x44, 0x88)
    RED_RGB = RGBColor(0xB0, 0x20, 0x20)
    INK_RGB = RGBColor(0x22, 0x28, 0x4A)
    SLATE_RGB = RGBColor(0x8A, 0x93, 0xB5)
    GOLD_RGB = RGBColor(0xC9, 0xA2, 0x4B)

    person = meta.get('name') or ''
    astrologer = meta.get('astrologer') or _brand().get('astrologo_default', '')
    hoy = fecha_es(_date.today().isoformat())

    def add_text_with_icons(par, text, bold=False, size=None, color=None):
        for seg, icon in _segments(text):
            if seg:
                r = par.add_run(seg)
                r.bold = bold
                if size:
                    r.font.size = size
                if color:
                    r.font.color.rgb = color
            if icon:
                try:
                    par.add_run(' ')
                    par.add_run().add_picture(icon, height=Inches(0.17))
                except Exception:
                    pass

    def ext_link(par, url, text, color='3A4488'):
        part = par.part
        r_id = part.relate_to(url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True)
        h = OxmlElement('w:hyperlink')
        h.set(qn('r:id'), r_id)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
        r.append(rPr)
        t = OxmlElement('w:t'); t.text = text
        r.append(t); h.append(r); par._p.append(h)

    def toc_field(par):
        """Campo TOC de Word: muestra títulos con números de página
        (clic derecho → Actualizar campos)."""
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'TOC \\o "1-2" \\h \\z \\u')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = "Índice — haz clic derecho y elige «Actualizar campos» para ver las páginas."
        r.append(t); fld.append(r)
        par._p.append(fld)

    doc = Document()
    # Letra base más grande y aireada (lectores adultos mayores)
    _normal = doc.styles['Normal']
    _normal.font.size = Pt(12.5)
    _normal.paragraph_format.space_after = Pt(10)
    _normal.paragraph_format.line_spacing = 1.25
    for s in doc.sections:
        s.top_margin = Inches(1); s.bottom_margin = Inches(1)
        s.left_margin = Inches(1); s.right_margin = Inches(1)

    def cover_line(text, size, color, bold=False, italic=False, font='Georgia',
                   space_after=6):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = font
        return p

    # Portada de marca (imagen) + página de datos
    cover_img = _cover_path(meta.get('chart_type'))
    if cover_img:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(0)
        try:
            # portada vertical: ajustar por altura para que quepa en la página
            cp.add_run().add_picture(cover_img, height=Inches(9.2))
        except Exception:
            pass
        doc.add_page_break()
        doc.add_paragraph()
    else:
        for _ in range(4):
            doc.add_paragraph()
    cover_line('Reporte Astrológico', 28, GOLD_RGB, bold=True)
    cover_line(meta.get('subtitle', 'de Carta Natal'), 15, SLATE_RGB, italic=True, space_after=22)
    if person:
        cover_line(person, 17, NAVY_RGB, bold=True, space_after=10)
    linea = (meta.get('time', '') + '  —  ' if meta.get('time') else '') + fecha_es(meta.get('date', ''))
    cover_line(linea, 12, INK_RGB)
    if meta.get('city'):
        cover_line(meta['city'], 12, INK_RGB, space_after=8)

    # Retorno solar: lugar y momento exacto del retorno
    if meta.get('chart_type') == 'solar_return' and meta.get('return_moment'):
        if meta.get('relocated') and meta.get('return_place'):
            cover_line('Relocalizado en: ' + meta['return_place'], 11, GOLD_RGB, italic=True)
        rmoment = meta['return_moment'] + (('  (' + meta['return_tz'] + ')') if meta.get('return_tz') else '')
        cover_line('El retorno solar exacto ocurre el', 10.5, SLATE_RGB)
        cover_line(rmoment, 12, GOLD_RGB, bold=True)
        if meta.get('return_ut'):
            cover_line(meta['return_ut'] + ' UT', 10, SLATE_RGB, space_after=20)

    # Combinada: segunda persona con sus datos
    pb = meta.get('person_b')
    if meta.get('chart_type') == 'combined' and pb:
        cover_line('en vínculo con', 10.5, SLATE_RGB, italic=True)
        if pb.get('name'):
            cover_line(pb['name'], 15, NAVY_RGB, bold=True)
        lb = (pb.get('time', '') + '  —  ' if pb.get('time') else '') + fecha_es(pb.get('date', ''))
        cover_line(lb, 12, INK_RGB)
        if pb.get('city'):
            cover_line(pb['city'], 12, INK_RGB, space_after=20)

    if astrologer:
        cover_line('Astrólogo: %s' % astrologer, 11.5, BLUE_RGB, italic=True)
    cover_line(hoy, 9.5, SLATE_RGB)
    cover_line('R E P O R T E   A S T R A L', 9, GOLD_RGB, space_after=0)
    doc.add_page_break()

    # Índice (campo de Word con números de página)
    hp = doc.add_heading(level=1)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("Índice de contenidos"); hr.font.color.rgb = NAVY_RGB
    toc_field(doc.add_paragraph())
    doc.add_page_break()

    def H2(text):
        hp = doc.add_heading(level=1)
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_with_icons(hp, text, color=NAVY_RGB)

    def H3(text):
        hp = doc.add_heading(level=2)
        add_text_with_icons(hp, text, color=BLUE_RGB)

    def body_p(text):
        bp = doc.add_paragraph()
        bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bp.paragraph_format.first_line_indent = Inches(0.18)
        add_text_with_icons(bp, text)

    def item_p(nm, tx, hl=False):
        ip = doc.add_paragraph()
        ip.paragraph_format.left_indent = Inches(0.2)
        add_text_with_icons(ip, nm, bold=True, color=RED_RGB if hl else INK_RGB)
        ip.add_run(" — " + tx)

    def flush_mand_docx(rows):
        if not rows:
            return
        tbl = doc.add_table(rows=len(rows), cols=2)
        tbl.allow_autofit = True
        for i, (label, frase, hl) in enumerate(rows):
            col = RED_RGB if hl else INK_RGB
            lc = tbl.cell(i, 0).paragraphs[0]
            lr = lc.add_run(label); lr.bold = True; lr.font.color.rgb = col
            lr.font.size = Pt(12)
            rc = tbl.cell(i, 1).paragraphs[0]
            rr = rc.add_run(frase); rr.font.color.rgb = col; rr.font.size = Pt(12)
        try:
            for row in tbl.rows:
                row.cells[0].width = Inches(1.9)
                row.cells[1].width = Inches(4.6)
        except Exception:
            pass

    if pre_blocks:
        mand_buf = []
        for blk in pre_blocks:
            if blk[0] == "mand":
                mand_buf.append((blk[1], blk[2], blk[3]))
                continue
            if mand_buf:
                flush_mand_docx(mand_buf); mand_buf = []
            if blk[0] == "h2":
                H2(blk[1])
            elif blk[0] == "h3":
                H3(blk[1])
            elif blk[0] == "p":
                body_p(blk[1])
            elif blk[0] == "item":
                item_p(blk[1], blk[2], blk[3])
        if mand_buf:
            flush_mand_docx(mand_buf); mand_buf = []

    if chart_png_bytes:
        caps = meta.get('img_captions')
        if chart_png2_bytes and caps:
            # Ambas ruedas en una página apaisada, lado a lado, para compararlas.
            from docx.enum.section import WD_ORIENT, WD_SECTION
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
            sec.orientation = WD_ORIENT.LANDSCAPE
            if sec.page_width < sec.page_height:
                sec.page_width, sec.page_height = sec.page_height, sec.page_width
            sec.left_margin = sec.right_margin = Inches(0.6)
            sec.top_margin = sec.bottom_margin = Inches(0.7)
            H2(meta.get('chart_heading', 'Tus dos cartas'))
            tbl = doc.add_table(rows=1, cols=2)
            for j, (img_bytes, cap) in enumerate(((chart_png_bytes, caps[0]),
                                                  (chart_png2_bytes, caps[1]))):
                cell = tbl.cell(0, j)
                cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(cap); cr.bold = True; cr.font.color.rgb = BLUE_RGB
                ip = cell.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    ip.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(4.7))
                except Exception:
                    pass
            sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
            sec2.orientation = WD_ORIENT.PORTRAIT
            if sec2.page_width > sec2.page_height:
                sec2.page_width, sec2.page_height = sec2.page_height, sec2.page_width
            sec2.left_margin = sec2.right_margin = Inches(1)
            sec2.top_margin = sec2.bottom_margin = Inches(1)
        else:
            H2(meta.get('chart_heading', 'Carta natal: el libreto de tu vida'))
            pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                pic.add_run().add_picture(io.BytesIO(chart_png_bytes), width=Inches(5.6))
            except Exception:
                pass

    if legend:
        H3("Leyenda de símbolos")
        for gname, rows in legend:
            gp = doc.add_paragraph()
            gr = gp.add_run(gname); gr.bold = True
            gr.font.color.rgb = BLUE_RGB
            ncols = 4
            nrows = (len(rows) + ncols - 1) // ncols
            tbl = doc.add_table(rows=nrows, cols=ncols)
            for i, (icon, label) in enumerate(rows):
                cell = tbl.cell(i // ncols, i % ncols)
                cp = cell.paragraphs[0]
                try:
                    cp.add_run().add_picture(icon, height=Inches(0.14))
                except Exception:
                    pass
                cp.add_run("  " + label).font.size = Pt(9.5)

    for sec_title, items in sections:
        H2(sec_title)
        for it_title, paras in items:
            H3(it_title)
            for p in paras:
                body_p(p)

    if glossary:
        for blk in glossary:
            if blk[0] == "h2":
                H2(blk[1])
            elif blk[0] == "h3":
                H3(blk[1])
            elif blk[0] == "p":
                body_p(blk[1])
            elif blk[0] == "item":
                item_p(blk[1], blk[2], False)

    # Cierre / palabras finales
    if closing:
        doc.add_page_break()
        for blk in closing:
            if blk[0] == "h2":
                H2(blk[1])
            elif blk[0] == "pc":
                cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(blk[1]); cr.italic = True
                cr.font.size = Pt(13); cr.font.color.rgb = NAVY_RGB

    # Antepenúltima: página de información (rueda educativa)
    info_img = _info_path()
    if info_img:
        doc.add_page_break()
        ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            ip.add_run().add_picture(info_img, width=Inches(6.9))
        except Exception:
            pass

    # Penúltima: aviso legal y de privacidad
    doc.add_page_break()
    H2("Aviso legal y de privacidad")
    for htitle, ptext in _DISCLAIMER:
        hp = doc.add_paragraph(); hr = hp.add_run(htitle)
        hr.bold = True; hr.font.color.rgb = NAVY_RGB; hr.font.size = Pt(12.5)
        dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        dr = dp.add_run(ptext); dr.font.size = Pt(10.5)

    brand = _brand()
    redes = brand.get('redes', [])
    if redes:
        doc.add_page_break()
        H2(brand.get('redes_titulo', 'Encuéntrame en'))
        if astrologer:
            cover_line(astrologer, 13, BLUE_RGB, italic=True)
        for r in redes:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ic = os.path.join(_ICONS_DIR, 'soc_' + r.get('icon', '') + '.png')
            if os.path.exists(ic):
                try:
                    p.add_run().add_picture(ic, height=Inches(0.15))
                    p.add_run('  ')
                except Exception:
                    pass
            rr = p.add_run(r.get('nombre', '') + ' — '); rr.bold = True
            rr.font.color.rgb = INK_RGB
            ext_link(p, r.get('url', '#'), r.get('texto', r.get('url', '')))

        # Nuestra app: Diario Astral
        _app = os.path.join(_COVERS_DIR, 'diario_astral.png')
        doc.add_paragraph()
        if os.path.exists(_app):
            ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                ap.add_run().add_picture(_app, height=Inches(0.7))
            except Exception:
                pass
        cover_line('Diario Astral: tu brújula cósmica personal, en español.',
                   13, NAVY_RGB, bold=True, space_after=2)
        cover_line('Descárgala en App Store y Google Play.',
                   11, SLATE_RGB, italic=True)

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ════════════════════════════════════════════════════════════════════════
#  ORQUESTACIÓN
# ════════════════════════════════════════════════════════════════════════

_SUBTITLE = {
    'natal': 'de Carta Natal', 'transit': 'de Tránsitos',
    'solar_return': 'de Retorno Solar', 'progressed': 'de Carta Progresada',
    'combined': 'de Carta Combinada', 'akashic': 'de Registros Akáshicos',
}
_CHART_HEADING = {
    'natal': 'Carta natal: el libreto de tu vida',
    'transit': 'Bi-rueda: tu carta natal (interior) y los tránsitos (exterior)',
    'solar_return': 'Tus dos cartas: natal y retorno solar',
    'progressed': 'Tus dos cartas: natal y progresada',
    'combined': 'La carta combinada del vínculo',
    'akashic': 'Tu carta natal y tu Casa XII',
}
_INTRO_TITLE = {
    'transit': 'Qué son los tránsitos',
    'solar_return': 'Qué es el retorno solar',
    'progressed': 'Qué es la carta progresada',
    'combined': 'Qué es la carta combinada',
}


def generate(data, name, fmt, chart_png, city="", astrologer="", chart_type="natal",
             chart_png2=None, name_b="", city_b=""):
    """Punto de entrada. Devuelve (bytes, filename, mimetype).

    `data` es el dict que devuelve astro para el tipo pedido:
      natal        → la carta natal
      transit      → {natal, transit, cross_aspects}
      solar_return → {natal, solar_return, year}
      progressed   → {natal, progressed, years}
      combined     → {a, b, combined}
    """
    ct = chart_type or "natal"

    if ct == "transit":
        chart = data['transit']
        base_meta = data.get('natal', chart)
        xa = [{'a': x['transit'], 'b': x['natal'], 'type': x['type'], 'orb': x['orb']}
              for x in data.get('cross_aspects', [])]
        _intro, sections = build_sections(chart, 'transit',
                                          house_key='natal_house', aspects=xa)
    elif ct == "solar_return":
        chart = data['solar_return']
        base_meta = data.get('natal', chart)
        _intro, sections = build_sections(chart, 'solar_return')
    elif ct == "progressed":
        chart = data['progressed']
        base_meta = data.get('natal', chart)
        _intro, sections = build_sections(chart, 'progressed')
    elif ct == "combined":
        chart = data['combined']
        base_meta = data.get('a', chart)
        _intro, sections = build_sections(chart, 'combined')
    elif ct == "akashic":
        # Registros Akáshicos: estudio de la Casa XII de la carta natal
        chart = data
        base_meta = data
        _ak_pre, sections = build_akashic(chart)
        _intro = ""
    else:
        ct = "natal"
        chart = data
        base_meta = data
        _intro, sections = build_sections(chart, 'natal')

    if ct == "akashic":
        pre_blocks = _ak_pre
    elif ct in ("natal", "solar_return", "combined", "progressed"):
        # Progresada usa el mismo preámbulo que natal/retorno (Sol, Luna, Asc,
        # cruz y elemento), adaptado.
        pre_blocks = build_preamble(chart, ct)
    elif ct == "transit":
        pre_blocks = [("h2", _INTRO_TITLE.get(ct, "Qué son los tránsitos"))]
        if _intro:
            pre_blocks.append(("p", _intro))
        pre_blocks.append(("h3", "El clima astral de estos tránsitos"))
        for para in build_transit_intro(data):
            pre_blocks.append(("p", para))
    else:
        pre_blocks = []
        if _INTRO_TITLE.get(ct):
            pre_blocks.append(("h2", _INTRO_TITLE[ct]))
        if _intro:
            pre_blocks.append(("p", _intro))

    subtitle = _SUBTITLE.get(ct, 'de Carta Natal')
    if ct == 'solar_return' and data.get('year'):
        subtitle += ' %s' % data['year']

    legend = build_legend()
    glossary = build_glossary()
    closing = build_closing(ct)
    person = (name or '').strip()
    meta = {
        'name': person,
        'city': (city or '').strip(),
        'astrologer': (astrologer or '').strip(),
        'date': base_meta.get('input', {}).get('date', ''),
        'time': base_meta.get('input', {}).get('time', ''),
        'subtitle': subtitle,
        'chart_heading': _CHART_HEADING.get(ct, 'Carta natal'),
        'chart_type': ct,
    }

    # Datos específicos por tipo para la portada
    if ct == 'solar_return':
        si = data.get('solar_return', {}).get('input', {})
        rl = si.get('return_local', '')          # 'YYYY-MM-DD HH:MM'
        if rl:
            parts = rl.split(' ')
            dp = parts[0]; tp = parts[1] if len(parts) > 1 else ''
            meta['return_moment'] = fecha_es(dp) + ((', ' + tp) if tp else '')
        else:
            meta['return_moment'] = ''
        meta['return_tz'] = si.get('return_tz', '')
        meta['return_ut'] = si.get('return_ut', '')
        meta['return_place'] = (data.get('reloc_city') or city or '').strip()
        meta['relocated'] = si.get('relocated', False)
        meta['img_captions'] = ['Carta natal', 'Retorno solar']
    elif ct == 'progressed':
        meta['img_captions'] = ['Carta natal', 'Carta progresada']
    elif ct == 'combined':
        bi = data.get('b', {}).get('input', {})
        meta['person_b'] = {
            'name': (name_b or '').strip(),
            'date': bi.get('date', ''), 'time': bi.get('time', ''),
            'city': (city_b or '').strip(),
        }

    # ── Datos para la lámina estética (nueva carta del reporte) ───────────
    if ct in ('natal', 'akashic', 'solar_return', 'progressed', 'combined'):
        meta['lamina_chart'] = chart
        _toc_lbl = {'natal': 'Carta natal', 'akashic': 'Registros akáshicos',
                    'solar_return': 'Retorno solar' + ((' %s' % data['year'])
                                                       if data.get('year') else ''),
                    'progressed': 'Carta progresada', 'combined': 'Carta combinada'}
        meta['lamina_toc'] = _toc_lbl.get(ct, 'Carta')

        def _coords(inp):
            try:
                return "%.2f°, %.2f°" % (float(inp.get('lat')), float(inp.get('lon')))
            except Exception:
                return ''

        def _clean(rows):
            return [(ic, t) for ic, t in rows if t]

        _SG = _rc.SIGN_NAME if _rc else ['']*12

        def _asc_row(chart_dict):
            a = chart_dict.get('angles', {}).get('asc', {}).get('lon')
            if a is None:
                return None
            return ('up', 'Asc %s %d°' % (_SG[int(a // 30) % 12], int(a % 30)))

        if ct in ('natal', 'akashic'):
            n_in = base_meta.get('input', {})
            rows = _clean([('cal', fecha_es(n_in.get('date', meta.get('date', '')))),
                           ('clock', n_in.get('time', '')), ('pin', (city or '').strip())])
            ar = _asc_row(chart)
            if ar:
                rows.append(ar)
            co = _coords(n_in)
            if co:
                rows.append(('globe', co))
            meta['boxes'] = [{'title': 'Nacimiento', 'rows': rows}]
        elif ct == 'solar_return':
            n_in = data.get('natal', {}).get('input', {})
            box1 = _clean([('cal', fecha_es(n_in.get('date', ''))),
                           ('clock', n_in.get('time', '')), ('pin', (city or '').strip())])
            rl = meta.get('return_moment', '')
            r2 = []
            if rl:
                parts = rl.split(',')
                r2.append(('cal', parts[0].strip()))
                if len(parts) > 1:
                    r2.append(('clock', parts[1].strip()))
            rplace = (meta.get('return_place') or city or '').strip()
            if rplace:
                r2.append(('pin', rplace))
            meta['boxes'] = [{'title': 'Nacimiento', 'rows': box1},
                             {'title': 'Retorno solar exacto', 'rows': _clean(r2)}]
            meta['lamina_natal'] = data.get('natal')
            _yr = data.get('year')
            meta['dual_captions'] = ['Carta natal',
                                     'Retorno solar' + ((' %s' % _yr) if _yr else '')]
            meta['dual_title'] = 'Tus dos cartas: natal y retorno solar'
            meta['dual_subs'] = [' · '.join(t for _, t in box1),
                                 ' · '.join(t for _, t in _clean(r2))]
        elif ct == 'progressed':
            n_in = data.get('natal', {}).get('input', {})
            box1 = _clean([('cal', fecha_es(n_in.get('date', ''))),
                           ('clock', n_in.get('time', '')), ('pin', (city or '').strip())])
            p_in = data.get('progressed', {}).get('input', {})
            import re as _re
            lbl = p_in.get('label') or ''
            m = _re.search(r'(\d{4}-\d{2}-\d{2})', lbl)
            r2 = [('cal', fecha_es(m.group(1)))] if m else [('up', lbl or 'Carta progresada')]
            meta['boxes'] = [{'title': 'Nacimiento', 'rows': box1},
                             {'title': 'Progresión', 'rows': _clean(r2)}]
            meta['lamina_natal'] = data.get('natal')
            meta['dual_captions'] = ['Carta natal', 'Carta progresada']
            meta['dual_title'] = 'Tus dos cartas: natal y progresada'
            meta['dual_subs'] = [' · '.join(t for _, t in box1),
                                 ' · '.join(t for _, t in _clean(r2))]
        elif ct == 'combined':
            a_in = data.get('a', {}).get('input', {})
            b_in = data.get('b', {}).get('input', {})
            pb = meta.get('person_b', {})
            nA = person or 'Persona 1'
            nB = (pb.get('name') or 'Persona 2')
            box1 = _clean([('cal', fecha_es(a_in.get('date', ''))),
                           ('clock', a_in.get('time', '')), ('pin', (city or '').strip())])
            box2 = _clean([('cal', fecha_es(b_in.get('date', ''))),
                           ('clock', b_in.get('time', '')), ('pin', (pb.get('city') or '').strip())])
            meta['boxes'] = [{'title': nA, 'rows': box1}, {'title': nB, 'rows': box2}]
            meta['lamina_name'] = nA + '  +  ' + nB

    png = _png_from_dataurl(chart_png)
    png2 = _png_from_dataurl(chart_png2)
    tsuf = '' if ct == 'natal' else ('_' + ct)
    safe = "Reporte_Astral" + (("_" + person.replace(' ', '_')) if person else "") + tsuf
    # El nombre del archivo va en una cabecera HTTP (solo ASCII): quita acentos
    # y cualquier carácter no seguro para no romper la respuesta.
    import unicodedata as _ud
    safe = ''.join(ch for ch in _ud.normalize('NFKD', safe) if ord(ch) < 128)
    safe = ''.join(ch if (ch.isalnum() or ch in ('_', '-')) else '_' for ch in safe) or 'Reporte_Astral'
    if fmt == 'docx':
        out = render_docx(sections, png, meta,
                          pre_blocks=pre_blocks, legend=legend, glossary=glossary,
                          chart_png2_bytes=png2, closing=closing)
        return out, safe + '.docx', \
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    out = render_pdf(sections, png, meta,
                     pre_blocks=pre_blocks, legend=legend, glossary=glossary,
                     chart_png2_bytes=png2, closing=closing)
    return out, safe + '.pdf', 'application/pdf'
